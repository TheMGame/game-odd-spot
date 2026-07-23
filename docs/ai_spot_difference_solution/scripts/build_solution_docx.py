from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "AI多地区找茬游戏完整方案.docx"
SOURCES = [
    ROOT / "docs/01_product_and_scope.md",
    ROOT / "docs/02_architecture_and_modules.md",
    ROOT / "docs/03_content_and_localization.md",
    ROOT / "docs/04_api_and_data.md",
    ROOT / "docs/05_deployment_and_operations.md",
    ROOT / "docs/06_development_roadmap.md",
    ROOT / "docs/07_state_machines_and_quality.md",
]


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=(100, 100, 100))
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_font(run, size=9, color=(100, 100, 100))


def add_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=9.5, color=(31, 78, 120))
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run)


def add_markdown(doc, path):
    in_code = False
    code_lines = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("\n".join(code_lines))
                set_font(run, name="Consolas", size=8.5)
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, line[2:])
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, line[3:])
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, line[4:])
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        else:
            p = doc.add_paragraph()
            add_inline(p, line)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, before, after, color in [
    ("Heading 1", 16, 18, 10, (46, 116, 181)),
    ("Heading 2", 13, 14, 7, (46, 116, 181)),
    ("Heading 3", 12, 10, 5, (31, 77, 120)),
]:
    style = styles[name]
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(*color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.text = "AI 多地区找茬游戏 · 产品与技术设计"
set_font(header.runs[0], size=9, color=(110, 110, 110))
add_page_number(section.footer.paragraphs[0])

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(72)
title.paragraph_format.space_after = Pt(12)
run = title.add_run("AI 多地区找茬游戏")
set_font(run, size=28, bold=True, color=(31, 77, 120))
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("完整产品、协议、数据与 Ubuntu 部署方案")
set_font(run, size=14, color=(80, 80, 80))
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(18)
run = meta.add_run("版本 2.0 · Godot + Go + MySQL + Redis · 原生 systemd 部署")
set_font(run, size=10, color=(110, 110, 110))
doc.add_page_break()

lead = doc.add_paragraph()
run = lead.add_run("文档权威性说明：")
set_font(run, bold=True)
add_inline(lead, "本 DOCX 用于评审阅读；发生冲突时，以分章 Markdown、JSON Schema、OpenAPI、SQL 和 deploy 目录为准。")

for index, source in enumerate(SOURCES):
    if index:
        doc.add_page_break()
    add_markdown(doc, source)

doc.add_page_break()
p = doc.add_paragraph(style="Heading 1")
add_inline(p, "附录：可执行设计资产")
for item in [
    "schemas/level.schema.json：客户端运行时关卡协议",
    "schemas/api.openapi.yaml：HTTP API 契约",
    "schemas/core_tables.sql：MySQL 8 核心表参考",
    "schemas/examples/：Schema 验收样例",
    "deploy/：Ubuntu systemd、Nginx、环境变量和发布脚本",
]:
    p = doc.add_paragraph(style="List Bullet")
    add_inline(p, item)

doc.save(OUTPUT)
print(OUTPUT)
